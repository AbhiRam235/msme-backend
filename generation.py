# backend/generation.py
import os
import uuid
import base64
from typing import Dict
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import pandas as pd
from templates import PROJECT_TEMPLATES

from finance import generate_financials


# Optional: OpenAI to generate textual sections
try:
    import openai
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False

OPENAI_MODEL = "gpt-4o"  # change to the model you use; the user may use GPT-5 via API

OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)

def classify_project(short_description: str) -> str:
    """
    Naive rule-based classification. Replace with a ML/NLP classifier or fine-tuned model.
    """
    d = short_description.lower()
    if "rice" in d or "agri" in d or "processing" in d or "farm" in d:
        return "agro_processing"
    if "ev" in d or "charging" in d or "electric vehicle" in d:
        return "ev_charging"
    return "default"

def fetch_external_data_stub(location: str) -> dict:
    """
    Placeholder for real external data fetch (markets, population, weather, land rates).
    Replace with calls to real APIs and caching.
    """
    return {
        "population_nearby": 120000,
        "avg_power_cost_per_kwh": 8.0,
        "land_rent_per_acre": 150000
    }

def generate_sections_with_openai(sections: list, project_payload: dict) -> Dict[str,str]:
    """
    If OpenAI available, generate textual content for each section.
    Otherwise return simple placeholders.
    """
    content = {}
    if OPENAI_AVAILABLE and os.getenv("OPENAI_API_KEY"):
        openai.api_key = os.getenv("OPENAI_API_KEY")
        for sec in sections:
            prompt = f"""Write a clear, professional section titled "{sec}" for a Detailed Project Report.
Project title: {project_payload.get('title')}
Short description: {project_payload.get('short_description')}
Location: {project_payload.get('location')}
Be concise but useful (~250-400 words). Include bullet points where useful."""
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role":"user","content":prompt}],
                max_tokens=700,
                temperature=0.2
            )
            text = resp.choices[0].message.content.strip()
            content[sec] = text
    else:
        # fallback placeholder text
        for sec in sections:
            content[sec] = f"[Auto-generated section: {sec}]\n\nProject: {project_payload.get('title')}\nDescription: {project_payload.get('short_description')}\n\n(Replace with real content or connect OpenAI API)"
    return content

def plot_financials(df: pd.DataFrame, out_png: Path):
    """
    Plot revenue and EBITDA.
    """
    plt.clf()
    ax = df.plot(kind="line", marker="o")
    ax.set_title("Revenue & EBITDA (projection)")
    ax.set_ylabel(df.columns[0])
    plt.tight_layout()
    plt.savefig(out_png, dpi=150)
    plt.close()

def create_docx(project_payload: dict, sections: dict, df: pd.DataFrame, meta: dict, out_path: Path):
    doc = Document()
    doc.add_heading(project_payload.get("title"), level=0)
    doc.add_paragraph(f"Location: {project_payload.get('location') or 'N/A'}")
    doc.add_paragraph(f"Currency: {meta.get('currency', 'INR')}")
    doc.add_page_break()

    for sec_title, body in sections.items():
        doc.add_heading(sec_title, level=1)
        # keep paragraphs short
        for para in body.split("\n\n"):
            doc.add_paragraph(para)

    # add financial table
    doc.add_page_break()
    doc.add_heading("Financial Projections (Summary)", level=1)
    t = doc.add_table(rows=1, cols=len(df.columns)+1)
    hdr = t.rows[0].cells
    hdr[0].text = "Year"
    for i, c in enumerate(df.columns, start=1):
        hdr[i].text = str(c)

    for idx, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(idx)
        for i, c in enumerate(df.columns, start=1):
            cells[i].text = f"{row[c]:,.2f}"

    doc.save(out_path)

def create_pdf_from_docx_and_chart(docx_path: Path, chart_png: Path, out_pdf: Path):
    """
    Simple PDF generator that places the chart and a note. For production,
    convert docx to PDF properly (LibreOffice headless, docx2pdf on Windows, or ReportLab full layout)
    """
    c = canvas.Canvas(str(out_pdf), pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width/2, height - 40, f"DPR - {docx_path.stem}")
    # embed chart
    c.drawImage(str(chart_png), 40, height - 380, width=520, preserveAspectRatio=True, mask='auto')
    c.setFont("Helvetica", 10)
    c.drawString(40, height - 410, "Note: Full textual DPR is included in the .docx. This PDF contains summary and charts.")
    c.showPage()
    c.save()

def generate_dpr_package(payload: dict) -> dict:
    """
    Orchestrator: classifies, fetches data, makes financials, generates docx and pdf,
    returns file paths.
    """
    uid = uuid.uuid4().hex[:8]
    project_type = classify_project(payload.get("short_description", ""))
    template = PROJECT_TEMPLATES.get(project_type, PROJECT_TEMPLATES["default"])
    external = fetch_external_data_stub(payload.get("location", ""))
    df, meta = generate_financials(payload)

    # generate textual sections (OpenAI or placeholder)
    sections = generate_sections_with_openai(template["sections"], payload)

    # outputs
    out_dir = OUTPUT_DIR / f"{uid}"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"{uid}_dpr.docx"
    chart_png = out_dir / f"{uid}_finance.png"
    pdf_path = out_dir / f"{uid}_summary.pdf"
    excel_path = out_dir / f"{uid}_financials.xlsx"

    # make chart
    plot_financials(df, chart_png)

    # save excel
    df.to_excel(excel_path)

    # make docx
    create_docx(payload, sections, df, meta, docx_path)

    # make pdf summary
    create_pdf_from_docx_and_chart(docx_path, chart_png, pdf_path)

    return {
        "uid": uid,
        "project_type": project_type,
        "template_name": template["name"],
        "docx": str(docx_path),
        "pdf_summary": str(pdf_path),
        "excel_financials": str(excel_path)
    }
